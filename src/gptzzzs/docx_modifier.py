import os
from pathlib import Path
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from gptzzzs.context_aware import change_text_contextual


def modify_docx(input_path, synonyms, adjectives, percent_synonyms=30, 
                ignore_quotes=True, percent_adjectives=60, use_pos_filtering=True):
    """
    Modify a DOCX file using context-aware text transformation while preserving formatting.
    
    :param input_path: Path to the input DOCX file
    :param synonyms: Dictionary of word -> list of synonyms
    :param adjectives: List of adjectives for emphasis
    :param percent_synonyms: Percentage of words to replace with synonyms
    :param ignore_quotes: Whether to ignore words in quotes
    :param percent_adjectives: Percentage of adjectives to emphasize
    :param use_pos_filtering: Whether to filter synonyms by part of speech
    :return: Path to the output file
    """
    
    # Load the document
    doc = Document(input_path)
    
    # Process each paragraph
    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            continue
            
        # Store original formatting
        original_alignment = paragraph.alignment
        original_style = paragraph.style
        
        # Process runs (text segments with consistent formatting)
        # We need to rebuild the paragraph while preserving run-level formatting
        runs_data = []
        full_text = ""
        
        for run in paragraph.runs:
            if run.text:
                runs_data.append({
                    'text': run.text,
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font_name': run.font.name,
                    'font_size': run.font.size,
                    'font_color': run.font.color.rgb if run.font.color.rgb else None,
                    'start_pos': len(full_text),
                    'end_pos': len(full_text) + len(run.text)
                })
                full_text += run.text
        
        # Apply context-aware transformation to the full paragraph text
        if full_text.strip():
            modified_text = change_text_contextual(
                full_text,
                synonyms=synonyms,
                adjectives=adjectives,
                percent_synonyms=percent_synonyms,
                ignore_quotes=ignore_quotes,
                percent_adjectives=percent_adjectives,
                use_pos_filtering=use_pos_filtering
            )
            
            # Clear the paragraph
            paragraph.clear()
            
            # Try to map formatting from original runs to modified text
            # This is a best-effort approach since word positions change
            if len(modified_text) > 0:
                # Simple approach: apply the first run's formatting to entire paragraph
                # This preserves basic formatting but not complex multi-format paragraphs
                new_run = paragraph.add_run(modified_text)
                
                if runs_data:
                    first_run = runs_data[0]
                    new_run.bold = first_run['bold']
                    new_run.italic = first_run['italic']
                    new_run.underline = first_run['underline']
                    
                    if first_run['font_name']:
                        new_run.font.name = first_run['font_name']
                    if first_run['font_size']:
                        new_run.font.size = first_run['font_size']
                    if first_run['font_color']:
                        new_run.font.color.rgb = first_run['font_color']
            
            # Restore paragraph-level formatting
            paragraph.alignment = original_alignment
            paragraph.style = original_style
    
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if not paragraph.text.strip():
                        continue
                    
                    original_alignment = paragraph.alignment
                    original_style = paragraph.style
                    
                    # Collect run data
                    runs_data = []
                    full_text = ""
                    
                    for run in paragraph.runs:
                        if run.text:
                            runs_data.append({
                                'text': run.text,
                                'bold': run.bold,
                                'italic': run.italic,
                                'underline': run.underline,
                                'font_name': run.font.name,
                                'font_size': run.font.size,
                                'font_color': run.font.color.rgb if run.font.color.rgb else None
                            })
                            full_text += run.text
                    
                    # Apply transformation
                    if full_text.strip():
                        modified_text = change_text_contextual(
                            full_text,
                            synonyms=synonyms,
                            adjectives=adjectives,
                            percent_synonyms=percent_synonyms,
                            ignore_quotes=ignore_quotes,
                            percent_adjectives=percent_adjectives,
                            use_pos_filtering=use_pos_filtering
                        )
                        
                        # Clear and recreate
                        paragraph.clear()
                        
                        if len(modified_text) > 0:
                            new_run = paragraph.add_run(modified_text)
                            
                            if runs_data:
                                first_run = runs_data[0]
                                new_run.bold = first_run['bold']
                                new_run.italic = first_run['italic']
                                new_run.underline = first_run['underline']
                                
                                if first_run['font_name']:
                                    new_run.font.name = first_run['font_name']
                                if first_run['font_size']:
                                    new_run.font.size = first_run['font_size']
                                if first_run['font_color']:
                                    new_run.font.color.rgb = first_run['font_color']
                        
                        paragraph.alignment = original_alignment
                        paragraph.style = original_style
    
    # Generate output filename
    input_file = Path(input_path)
    output_filename = f"{input_file.stem}_edited{input_file.suffix}"
    output_path = input_file.parent / output_filename
    
    # Save the modified document
    doc.save(str(output_path))
    
    return str(output_path)


def modify_docx_from_gptzzzs(input_path, gptzzzs_instance, synonym_list="finnlp",
                             percent_synonyms=30, ignore_quotes=True,
                             percent_adjectives=60, common_words=False,
                             adjective_list="normal", use_pos_filtering=True):
    """
    Convenience function to modify a DOCX using a Gptzzzs instance.
    
    :param input_path: Path to the input DOCX file
    :param gptzzzs_instance: Instance of Gptzzzs class
    :param synonym_list: the list of synonyms to use (finnlp, zaibacu, custom)
    :param percent_synonyms: the percentage of words to change
    :param ignore_quotes: whether to ignore words in quotes
    :param percent_adjectives: the percentage of adjectives to emphasize
    :param common_words: whether to only use common words
    :param adjective_list: the list of adjectives to use (normal, custom)
    :param use_pos_filtering: whether to filter synonyms by part of speech
    :return: Path to the output file
    """
    
    # Load synonyms and adjectives from the instance
    if synonym_list == "finnlp":
        synonyms = gptzzzs_instance.finnlp_together
    elif synonym_list == "zaibacu":
        synonyms = gptzzzs_instance.zaibacu
    elif synonym_list == "custom" and gptzzzs_instance.custom_synonyms is None:
        raise ValueError("Custom synonyms not set")
    elif synonym_list == "custom":
        synonyms = gptzzzs_instance.custom_synonyms
    else:
        raise ValueError("Invalid synonym list")

    if adjective_list == "normal":
        adjectives = gptzzzs_instance.adjectives
    elif adjective_list == "custom" and gptzzzs_instance.custom_adjectives is None:
        raise ValueError("Custom adjectives not set")
    elif adjective_list == "custom":
        adjectives = gptzzzs_instance.custom_adjectives
    else:
        raise ValueError("Invalid adjective list")

    if common_words:
        synonyms = {word: synonyms[word] for word in synonyms if word in gptzzzs_instance.common_words}
        synonyms = {word: syn_list for word, syn_list in synonyms.items() if word in gptzzzs_instance.common_words and syn_list}
    
    return modify_docx(
        input_path=input_path,
        synonyms=synonyms,
        adjectives=adjectives,
        percent_synonyms=percent_synonyms,
        ignore_quotes=ignore_quotes,
        percent_adjectives=percent_adjectives,
        use_pos_filtering=use_pos_filtering
    )




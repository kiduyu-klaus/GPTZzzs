from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from gptzzzs.ollama_humanize import humanize_with_ollama


def process_docx(input_path, ollama_model="llama2", ollama_url="http://localhost:11434", 
                 temperature=0.7, max_tokens=2000, preserve_formatting=True):
    """
    Reads a docx file, humanizes the text using Ollama while preserving formatting,
    and saves to a new file with '_edited' suffix.
    
    :param input_path: path to the input docx file
    :param ollama_model: the Ollama model to use
    :param ollama_url: the URL of the Ollama API
    :param temperature: controls randomness (0.0-1.0)
    :param max_tokens: maximum number of tokens per request
    :param preserve_formatting: whether to preserve text formatting
    :return: path to the output file
    """
    
    # Validate input file
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"Input file not found: {input_path}")
    
    if not input_path.endswith('.docx'):
        raise ValueError("Input file must be a .docx file")
    
    # Generate output path
    directory = os.path.dirname(input_path)
    filename = os.path.basename(input_path)
    name, ext = os.path.splitext(filename)
    output_path = os.path.join(directory, f"{name}_edited{ext}")
    
    print(f"Reading document: {input_path}")
    
    # Load the document
    doc = Document(input_path)
    
    # Process each paragraph
    total_paragraphs = len(doc.paragraphs)
    for i, paragraph in enumerate(doc.paragraphs):
        # Skip empty paragraphs
        if not paragraph.text.strip():
            continue
        
        print(f"Processing paragraph {i + 1}/{total_paragraphs}...")
        
        original_text = paragraph.text
        
        try:
            # Humanize the text
            humanized_text = humanize_with_ollama(
                original_text,
                model=ollama_model,
                ollama_url=ollama_url,
                temperature=temperature,
                max_tokens=max_tokens
            )
            
            if preserve_formatting:
                # Store original formatting for each run
                run_formats = []
                for run in paragraph.runs:
                    run_formats.append({
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline,
                        'font_name': run.font.name,
                        'font_size': run.font.size,
                        'font_color': run.font.color.rgb if run.font.color.rgb else None
                    })
                
                # Store paragraph formatting
                paragraph_format = {
                    'alignment': paragraph.alignment,
                    'left_indent': paragraph.paragraph_format.left_indent,
                    'right_indent': paragraph.paragraph_format.right_indent,
                    'first_line_indent': paragraph.paragraph_format.first_line_indent,
                    'space_before': paragraph.paragraph_format.space_before,
                    'space_after': paragraph.paragraph_format.space_after,
                    'line_spacing': paragraph.paragraph_format.line_spacing
                }
                
                # Clear all runs in the paragraph
                for run in paragraph.runs:
                    run.text = ''
                
                # Add the humanized text to the first run (or create new one)
                if paragraph.runs:
                    paragraph.runs[0].text = humanized_text
                    # Apply the first run's original formatting
                    if run_formats:
                        apply_run_format(paragraph.runs[0], run_formats[0])
                else:
                    run = paragraph.add_run(humanized_text)
                    if run_formats:
                        apply_run_format(run, run_formats[0])
                
                # Restore paragraph formatting
                paragraph.alignment = paragraph_format['alignment']
                paragraph.paragraph_format.left_indent = paragraph_format['left_indent']
                paragraph.paragraph_format.right_indent = paragraph_format['right_indent']
                paragraph.paragraph_format.first_line_indent = paragraph_format['first_line_indent']
                paragraph.paragraph_format.space_before = paragraph_format['space_before']
                paragraph.paragraph_format.space_after = paragraph_format['space_after']
                paragraph.paragraph_format.line_spacing = paragraph_format['line_spacing']
            else:
                # Simple replacement without preserving formatting
                paragraph.text = humanized_text
                
        except Exception as e:
            print(f"Warning: Could not process paragraph {i + 1}: {str(e)}")
            print("Keeping original text for this paragraph.")
            continue
    
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if not paragraph.text.strip():
                        continue
                    
                    original_text = paragraph.text
                    
                    try:
                        humanized_text = humanize_with_ollama(
                            original_text,
                            model=ollama_model,
                            ollama_url=ollama_url,
                            temperature=temperature,
                            max_tokens=max_tokens
                        )
                        
                        if preserve_formatting and paragraph.runs:
                            # Store formatting of first run
                            run_format = {
                                'bold': paragraph.runs[0].bold,
                                'italic': paragraph.runs[0].italic,
                                'underline': paragraph.runs[0].underline,
                                'font_name': paragraph.runs[0].font.name,
                                'font_size': paragraph.runs[0].font.size,
                                'font_color': paragraph.runs[0].font.color.rgb if paragraph.runs[0].font.color.rgb else None
                            }
                            
                            # Clear and replace
                            for run in paragraph.runs:
                                run.text = ''
                            paragraph.runs[0].text = humanized_text
                            apply_run_format(paragraph.runs[0], run_format)
                        else:
                            paragraph.text = humanized_text
                            
                    except Exception as e:
                        print(f"Warning: Could not process table cell: {str(e)}")
                        continue
    
    # Save the modified document
    print(f"Saving edited document: {output_path}")
    doc.save(output_path)
    print("Done!")
    
    return output_path


def apply_run_format(run, format_dict):
    """
    Applies formatting to a run based on a format dictionary.
    
    :param run: the run to format
    :param format_dict: dictionary containing formatting info
    """
    if format_dict['bold'] is not None:
        run.bold = format_dict['bold']
    if format_dict['italic'] is not None:
        run.italic = format_dict['italic']
    if format_dict['underline'] is not None:
        run.underline = format_dict['underline']
    if format_dict['font_name']:
        run.font.name = format_dict['font_name']
    if format_dict['font_size']:
        run.font.size = format_dict['font_size']
    if format_dict['font_color']:
        run.font.color.rgb = format_dict['font_color']


def batch_process_docx(directory, file_pattern="*.docx", ollama_model="llama2", 
                       ollama_url="http://localhost:11434", temperature=0.7, 
                       max_tokens=2000, preserve_formatting=True):
    """
    Processes multiple docx files in a directory.
    
    :param directory: directory containing docx files
    :param file_pattern: pattern to match files (default: *.docx)
    :param ollama_model: the Ollama model to use
    :param ollama_url: the URL of the Ollama API
    :param temperature: controls randomness
    :param max_tokens: maximum tokens per request
    :param preserve_formatting: whether to preserve formatting
    :return: list of output file paths
    """
    import glob
    
    if not os.path.exists(directory):
        raise FileNotFoundError(f"Directory not found: {directory}")
    
    # Find all matching docx files (excluding _edited files)
    pattern = os.path.join(directory, file_pattern)
    files = [f for f in glob.glob(pattern) if not f.endswith('_edited.docx')]
    
    if not files:
        print(f"No docx files found in {directory}")
        return []
    
    output_files = []
    
    for i, input_file in enumerate(files):
        print(f"\n{'=' * 80}")
        print(f"Processing file {i + 1}/{len(files)}: {os.path.basename(input_file)}")
        print('=' * 80)
        
        try:
            output_file = process_docx(
                input_file,
                ollama_model=ollama_model,
                ollama_url=ollama_url,
                temperature=temperature,
                max_tokens=max_tokens,
                preserve_formatting=preserve_formatting
            )
            output_files.append(output_file)
        except Exception as e:
            print(f"Error processing {input_file}: {str(e)}")
            continue
    
    print(f"\n{'=' * 80}")
    print(f"Batch processing complete! Processed {len(output_files)}/{len(files)} files.")
    print('=' * 80)
    
    return output_files


def process_docx_with_progress(input_path, ollama_model="llama2", ollama_url="http://localhost:11434",
                               temperature=0.7, max_tokens=2000, preserve_formatting=True,
                               progress_callback=None):
    """
    Same as process_docx but with progress callback support.
    
    :param input_path: path to the input docx file
    :param ollama_model: the Ollama model to use
    :param ollama_url: the URL of the Ollama API
    :param temperature: controls randomness
    :param max_tokens: maximum tokens per request
    :param preserve_formatting: whether to preserve formatting
    :param progress_callback: callback function(current, total, message)
    :return: path to the output file
    """
    
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"Input file not found: {input_path}")
    
    if not input_path.endswith('.docx'):
        raise ValueError("Input file must be a .docx file")
    
    directory = os.path.dirname(input_path)
    filename = os.path.basename(input_path)
    name, ext = os.path.splitext(filename)
    output_path = os.path.join(directory, f"{name}_edited{ext}")
    
    doc = Document(input_path)
    total_paragraphs = len([p for p in doc.paragraphs if p.text.strip()])
    current = 0
    
    if progress_callback:
        progress_callback(0, total_paragraphs, "Starting processing...")
    
    for i, paragraph in enumerate(doc.paragraphs):
        if not paragraph.text.strip():
            continue
        
        current += 1
        if progress_callback:
            progress_callback(current, total_paragraphs, f"Processing paragraph {current}/{total_paragraphs}")
        
        original_text = paragraph.text
        
        try:
            humanized_text = humanize_with_ollama(
                original_text,
                model=ollama_model,
                ollama_url=ollama_url,
                temperature=temperature,
                max_tokens=max_tokens
            )
            
            if preserve_formatting:
                run_formats = []
                for run in paragraph.runs:
                    run_formats.append({
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline,
                        'font_name': run.font.name,
                        'font_size': run.font.size,
                        'font_color': run.font.color.rgb if run.font.color.rgb else None
                    })
                
                paragraph_format = {
                    'alignment': paragraph.alignment,
                    'left_indent': paragraph.paragraph_format.left_indent,
                    'right_indent': paragraph.paragraph_format.right_indent,
                    'first_line_indent': paragraph.paragraph_format.first_line_indent,
                    'space_before': paragraph.paragraph_format.space_before,
                    'space_after': paragraph.paragraph_format.space_after,
                    'line_spacing': paragraph.paragraph_format.line_spacing
                }
                
                for run in paragraph.runs:
                    run.text = ''
                
                if paragraph.runs:
                    paragraph.runs[0].text = humanized_text
                    if run_formats:
                        apply_run_format(paragraph.runs[0], run_formats[0])
                else:
                    run = paragraph.add_run(humanized_text)
                    if run_formats:
                        apply_run_format(run, run_formats[0])
                
                paragraph.alignment = paragraph_format['alignment']
                paragraph.paragraph_format.left_indent = paragraph_format['left_indent']
                paragraph.paragraph_format.right_indent = paragraph_format['right_indent']
                paragraph.paragraph_format.first_line_indent = paragraph_format['first_line_indent']
                paragraph.paragraph_format.space_before = paragraph_format['space_before']
                paragraph.paragraph_format.space_after = paragraph_format['space_after']
                paragraph.paragraph_format.line_spacing = paragraph_format['line_spacing']
            else:
                paragraph.text = humanized_text
                
        except Exception as e:
            if progress_callback:
                progress_callback(current, total_paragraphs, f"Error in paragraph {current}: {str(e)}")
            continue
    
    if progress_callback:
        progress_callback(total_paragraphs, total_paragraphs, "Saving document...")
    
    doc.save(output_path)
    
    if progress_callback:
        progress_callback(total_paragraphs, total_paragraphs, "Done!")
    
    return output_path